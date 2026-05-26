from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "showcase"
ASSET_DIR = OUT_DIR / "assets"
ARCHITECTURE = ASSET_DIR / "phishguard-system-architecture.png"
SCREENSHOTS = Path(r"C:\Users\BAO MINH\Pictures\TechFest")
OUTPUT = OUT_DIR / "PhishGuard_SOC_System_Presentation_Report.docx"

NAVY = "12315B"
BLUE = "1767AE"
TEAL = "087F83"
LIGHT_BLUE = "EAF3FA"
LIGHT_TEAL = "E8F5F4"
GOLD = "8A5A00"
LIGHT_GOLD = "FFF4DC"
RED = "9B1C1C"
LIGHT_RED = "FCEAEA"
GRAY = "5D6776"
LIGHT_GRAY = "F2F4F7"
BORDER = "D6DEE8"
INK = "1F2937"


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    width = table_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        table_pr.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for col_width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(col_width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def border_table(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def cell_text(cell, text, bold=False, color=INK, size=9.2, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.08
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    set_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    border_table(table)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, header in enumerate(headers):
        shade_cell(table.cell(0, i), LIGHT_GRAY)
        cell_text(table.cell(0, i), header, bold=True, color=NAVY, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=9)
            set_cell_margins(cells[i])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_para(doc, text="", style=None, bold=False, color=INK, size=None, italic=False, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=9, color=GRAY, italic=True)


def add_figure(doc, path, caption, width=6.35, alt_text=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text or caption)
    doc_pr.set("title", caption)
    add_caption(doc, caption)


def add_callout(doc, title, body, tone="blue"):
    fills = {"blue": LIGHT_BLUE, "teal": LIGHT_TEAL, "gold": LIGHT_GOLD, "red": LIGHT_RED}
    heading_colors = {"blue": BLUE, "teal": TEAL, "gold": GOLD, "red": RED}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    border_table(table, heading_colors[tone])
    cell = table.cell(0, 0)
    shade_cell(cell, fills[tone])
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, color=heading_colors[tone], bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(body)
    set_font(r, size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_labeled_para(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(label + " ")
    set_font(r, size=11, color=NAVY, bold=True)
    r = p.add_run(body)
    set_font(r, size=11, color=INK)


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(9)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(TEAL)
    subtitle.paragraph_format.space_after = Pt(18)

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)


def set_page_layout(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PHISHGUARD SOC  |  LIVE GMAIL PROTOTYPE")
    set_font(r, size=8.5, color=GRAY, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("System Presentation Report  |  University Showcase Candidate  |  Page ")
    set_font(r, size=8, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def screenshot(name):
    return SCREENSHOTS / name


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if not ARCHITECTURE.exists():
        raise FileNotFoundError(f"Architecture asset missing: {ARCHITECTURE}")
    doc = Document()
    set_styles(doc)
    set_page_layout(doc)
    doc.core_properties.title = "PhishGuard SOC: Live Gmail Explainable Phishing Detection Dashboard"
    doc.core_properties.subject = "System presentation and operating guide"
    doc.core_properties.author = "Student Project Author"
    doc.core_properties.keywords = "PhishGuard, Gmail, phishing detection, SOC dashboard, explainable AI"

    # Cover page - editorial report opening.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SYSTEM PRESENTATION REPORT")
    set_font(r, size=11, color=TEAL, bold=True)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PhishGuard SOC")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Live Gmail Explainable Phishing Detection and Model-Governance Dashboard")
    add_para(
        doc,
        "A practical cybersecurity software prototype that scans Gmail on demand, explains risk decisions, "
        "supports reversible containment, and keeps analysts responsible for validation.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        color=GRAY,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(22)
    add_table(
        doc,
        ["PRODUCT STATUS", "PRIMARY USER", "OPERATING MODE"],
        [["Implemented live prototype", "SOC analyst / model owner", "Manual Gmail sync; human review"]],
        [2500, 3000, 3860],
    )
    add_callout(
        doc,
        "Positioning",
        "This is a Gmail prototype for an analyst-centred security workflow. It is not a production replacement "
        "for an enterprise secure email gateway, and it does not claim live accuracy without validated operational labels.",
        "gold",
    )
    add_para(doc, "Prepared for university showcase review  |  25 May 2026", color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_para(doc, "1. Executive Summary", style="Heading 1")
    add_para(
        doc,
        "PhishGuard SOC is a full-stack security operations prototype developed from an explainable phishing "
        "detection foundation into an interactive Gmail workflow. I built the system to address a practical "
        "problem: a classifier alone is not enough for a business security team. Analysts need to understand "
        "what was scanned, why a message was treated as risky, what action occurred, and how mistakes are reviewed.",
    )
    add_para(
        doc,
        "The implemented product connects to Gmail when locally configured, synchronises mailbox messages only "
        "when the analyst presses Sync Gmail, stores operational evidence in SQLite, classifies messages with a "
        "Random Forest teacher model, and uses an EBM surrogate for model-level explanation. It presents local "
        "evidence for a single case and global model-behaviour evidence for governance use.",
    )
    add_callout(
        doc,
        "What the screenshots prove",
        "The supplied application captures show a working analyst dashboard with 29 scanned Gmail messages, "
        "six phishing classifications, one quarantined email, four confirmed feedback cases, validated model "
        "readiness, successful manual sync status, local feature contributions, and separately labelled benchmark fidelity metrics.",
        "teal",
    )
    add_para(doc, "Implemented Capability Set", style="Heading 2")
    add_table(
        doc,
        ["CAPABILITY", "IMPLEMENTED BEHAVIOUR", "USER VALUE"],
        [
            ["Gmail ingestion", "Manual Sync Gmail action with status reporting and deduplication", "Analyst controls when mailbox data is processed"],
            ["Detection", "Risk prediction, confidence, risk level and action state", "Prioritises investigation work"],
            ["Local explanation", "Human-readable reasoning plus raw feature contributions", "Supports case-level decisions"],
            ["Global explanation", "EBM-based importance and fidelity benchmark view", "Supports model governance"],
            ["Containment", "Gmail label-based quarantine and reversible release", "Avoids permanent message deletion"],
            ["Feedback", "Analyst-confirmed review flow and monitoring counts", "Creates controlled learning evidence"],
        ],
        [1900, 4450, 3010],
    )

    doc.add_page_break()
    add_para(doc, "2. System Architecture", style="Heading 1")
    add_para(
        doc,
        "The solution is organised as a clear operational chain: Gmail supplies messages; the FastAPI backend "
        "coordinates sync and actions; the ML and explainability service produces evidence; SQLite stores case, "
        "prediction, explanation and audit state; and the React dashboard exposes the analyst workflow.",
    )
    add_figure(
        doc,
        ARCHITECTURE,
        "Figure 1. Implemented high-level architecture and analyst feedback loop.",
        width=6.45,
        alt_text="Architecture showing Gmail, FastAPI, ML and explainability, SQLite, React dashboard, and analyst feedback workflow.",
    )
    add_callout(
        doc,
        "Design principle",
        "Predictions support investigation; they do not prove malicious intent. Quarantine and release remain reversible analyst-facing actions.",
        "blue",
    )
    add_para(doc, "Runtime Components", style="Heading 2")
    add_table(
        doc,
        ["COMPONENT", "IMPLEMENTATION", "RESPONSIBILITY"],
        [
            ["Mailbox source", "Gmail API provider", "Fetch inbox messages and apply review-label actions"],
            ["API/backend", "Python FastAPI", "Validation, orchestration, monitoring and audit events"],
            ["ML/XAI", "Random Forest + EBM", "Prediction, readiness validation and explanations"],
            ["Persistence", "SQLite / SQLAlchemy", "Emails, latest predictions, feedback, sync and audit state"],
            ["Dashboard", "React + Vite + Recharts", "SOC workflow views and analyst controls"],
        ],
        [1900, 2540, 4920],
    )

    doc.add_page_break()
    add_para(doc, "3. Operational Workflow", style="Heading 1")
    add_para(
        doc,
        "The workflow is designed around controlled operation rather than invisible automation. In this prototype, "
        "Gmail synchronisation is initiated by an analyst and any containment is reversible.",
    )
    workflow = [
        ("Step 1 - Connect safely", "The operator configures Gmail OAuth credential and token paths through local environment settings. Secrets are excluded from source control."),
        ("Step 2 - Trigger mailbox sync", "The analyst presses Sync Gmail. The backend fetches the selected inbox messages, preserves URLs and metadata, and skips unchanged messages already processed under the current model state."),
        ("Step 3 - Validate and predict", "Before a decision is trusted, the runtime checks the teacher model, EBM surrogate, reconstructed fitted preprocessor, and exact 292-column feature order."),
        ("Step 4 - Record evidence", "SQLite stores the email record, current prediction, explanation snapshot, model version, sync state and relevant audit event."),
        ("Step 5 - Investigate and act", "The analyst reviews prediction and explanation evidence, then may confirm phishing, report an error, quarantine or release a message."),
        ("Step 6 - Monitor outcomes", "The dashboard separates live operational counts from research benchmark fidelity and reflects analyst-confirmed feedback without automatic retraining."),
    ]
    for title, body in workflow:
        add_labeled_para(doc, title + ":", body)
    add_para(doc, "Risk and Action Policy", style="Heading 2")
    add_table(
        doc,
        ["RISK LEVEL", "PROTOTYPE ACTION", "SAFETY INTERPRETATION"],
        [
            ["Low", "Allow / retain in inbox", "Visible for audit and review"],
            ["Medium", "Flag for analyst attention", "Does not require automatic removal"],
            ["High / Critical", "Reversible quarantine when prediction is trusted", "Add review label and remove INBOX; never delete"],
            ["Model error / unsafe pipeline", "Needs review", "No silent safe classification or trusted containment"],
        ],
        [1700, 3300, 4360],
    )

    doc.add_page_break()
    add_para(doc, "4. How to Run and Use the Prototype", style="Heading 1")
    add_para(doc, "Local Start-Up", style="Heading 2")
    add_labeled_para(doc, "Backend:", "From the project root, run start_backend.bat. The FastAPI service starts at http://127.0.0.1:8000 and reports Gmail/model readiness warnings without printing credential contents.")
    add_labeled_para(doc, "Frontend:", "From the project root, run start_frontend.bat. The Vite dashboard opens at http://localhost:5173.")
    add_labeled_para(doc, "Gmail setup:", "Use a local .env based on .env.example with paths to OAuth files stored outside Git tracking. The dashboard remains usable for review only when Gmail is not configured.")
    add_para(doc, "Analyst Use Procedure", style="Heading 2")
    use_steps = [
        ("1. Verify readiness", "Open Overview, Settings or Model Monitoring and confirm the Model Readiness Check reports Pipeline validated before trusting quarantine policy actions."),
        ("2. Synchronise", "Press Sync Gmail in the lower-left navigation control. Check Last Sync Status for scanned, skipped and failed counts."),
        ("3. Triage", "Open Detection Queue, search or filter cases and choose a message requiring investigation."),
        ("4. Explain", "Use Email Investigation to inspect local explanation text, raw contributions, confidence and model/action status."),
        ("5. Contain or release", "Use Quarantine or investigation controls; Gmail containment is reversible."),
        ("6. Review feedback", "Confirm reported mistakes in Feedback Review, then consult Model Monitoring for validated operational counts."),
    ]
    add_table(doc, ["ACTION", "HOW TO PERFORM IT"], use_steps, [1900, 7460])
    add_callout(
        doc,
        "Showcase privacy note",
        "The screenshots embedded in this report display real mailbox identifiers and message subjects. Before publishing the report publicly or submitting it outside an assessed/private context, redact personal email addresses and any sensitive subject content.",
        "red",
    )

    doc.add_page_break()
    add_para(doc, "5. Dashboard Evidence: Operational Overview", style="Heading 1")
    add_para(
        doc,
        "The Overview page establishes that this is an operational dashboard rather than a static prediction demo. "
        "It displays database-derived counts, model readiness, manual Gmail sync status, trend visualisation and risk distribution.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 212946.png"),
        "Figure 2. Overview page: operational metrics, validated pipeline status and successful Gmail sync status.",
        alt_text="Dashboard overview with operational metric cards, model readiness validation and last sync panel.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 212959.png"),
        "Figure 3. Overview analytics: detection trend, risk distribution and risky sender-domain aggregation.",
        alt_text="Dashboard analytics showing detection trend chart, risk distribution donut chart and sender domain table.",
    )
    add_para(doc, "Evidence Observed", style="Heading 2")
    add_table(
        doc,
        ["OBSERVATION", "VISIBLE EVIDENCE", "MEANING"],
        [
            ["Operational records", "29 scanned, 6 phishing, 1 quarantined", "Live Gmail workflow state is surfaced"],
            ["Pipeline readiness", "Teacher, surrogate and preprocessor ready; feature order recovered", "The UI communicates model safety state"],
            ["Manual sync state", "Success; 0 scanned, 25 skipped, 0 failed in latest run", "Deduplication prevents unnecessary reprocessing"],
            ["Trend and risk views", "Line/area and distribution charts", "Analysts can identify activity patterns quickly"],
        ],
        [1900, 3100, 4360],
    )

    doc.add_page_break()
    add_para(doc, "6. Dashboard Evidence: Queue and Containment", style="Heading 1")
    add_para(
        doc,
        "The Detection Queue converts scanned Gmail messages into cases with prediction, confidence, risk level, "
        "review state and action status. The Quarantine page provides a compact reversible containment queue.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213021.png"),
        "Figure 4. Detection Queue: Gmail cases with filtering, prediction, confidence, risk and action state.",
        alt_text="Detection queue listing Gmail message cases, predictions, confidence, risk level and status.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213026.png"),
        "Figure 5. Quarantine Review Queue: a contained message with a visible Release action.",
        alt_text="Quarantine queue showing one contained case and a release button.",
    )
    add_callout(
        doc,
        "Operational significance",
        "The visible Release action is important in a business setting: an incorrectly contained legitimate email can be returned to the inbox, while the case history remains available for review and audit.",
        "teal",
    )

    doc.add_page_break()
    add_para(doc, "7. Dashboard Evidence: Individual Email Investigation", style="Heading 1")
    add_para(
        doc,
        "The investigation workspace is the central analyst screen. It links message context to the active model "
        "version, explanation snapshot, prediction status and actions, allowing a human decision to be made with evidence.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213035.png"),
        "Figure 6. Investigation summary: case metadata, model decision and local explanation narrative.",
        alt_text="Investigation view showing email case metadata, prediction card and explanation narrative.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213044.png"),
        "Figure 7. Technical evidence and analyst actions: contribution graph, factor list and reversible controls.",
        alt_text="Investigation view showing feature contribution bars, explanation factors and analyst action buttons.",
    )
    add_para(doc, "Explanation Quality", style="Heading 2")
    add_para(
        doc,
        "The user-facing wording intentionally states that a feature influenced the model and does not prove malicious "
        "intent. This is appropriate for a security decision support tool: the interface provides both a readable summary "
        "for non-specialists and raw contribution values for technically trained analysts.",
    )
    add_table(
        doc,
        ["VISIBLE FEATURE", "DASHBOARD INTERPRETATION"],
        [
            ["Local narrative", "Explains model influence in cautious, human-readable language"],
            ["Contribution chart", "Shows direction and relative strength of technical signals"],
            ["Model version and snapshot", "Preserves which model/explanation supported the case decision"],
            ["Analyst action panel", "Turns evidence into controlled review, quarantine or release activity"],
        ],
        [2600, 6760],
    )

    doc.add_page_break()
    add_para(doc, "8. Dashboard Evidence: Global Explanation and Governance", style="Heading 1")
    add_para(
        doc,
        "The Global Explanation view explains model behaviour across the benchmark context rather than claiming to "
        "explain the correctness of any individual live message. The page clearly labels its fidelity results as research benchmark evidence.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213049.png"),
        "Figure 8. Global explanation view: benchmark notice, model pair and fidelity measures.",
        alt_text="Global explanation page showing benchmark warning, model pair and fidelity metric cards.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213055.png"),
        "Figure 9. Global feature importance: EBM-compatible model-behaviour view and failure-pattern area.",
        alt_text="Feature importance chart with ranked model factors and failure pattern summary panel.",
    )
    add_para(doc, "What the Benchmark Cards Mean", style="Heading 2")
    add_table(
        doc,
        ["METRIC SHOWN", "DISPLAYED VALUE", "CORRECT INTERPRETATION"],
        [
            ["Accuracy Fidelity", "92.6%", "Agreement between surrogate and teacher on benchmark cases"],
            ["F1 Fidelity", "92.8%", "Benchmark fidelity focused on classification balance"],
            ["Error Fidelity", "76.7%", "How reliably the surrogate follows teacher behaviour on error cases"],
        ],
        [2400, 1800, 5160],
    )
    add_callout(
        doc,
        "Truthful reporting boundary",
        "These fidelity cards are not live Gmail detection accuracy. Live operational performance requires sufficient analyst-confirmed labels, which the dashboard explicitly tracks separately.",
        "gold",
    )

    doc.add_page_break()
    add_para(doc, "9. Dashboard Evidence: Feedback and Model Monitoring", style="Heading 1")
    add_para(
        doc,
        "A cybersecurity classifier becomes more useful when mistakes can be reviewed without automatically changing "
        "the model. PhishGuard records feedback, requires analyst confirmation, and reflects validated results in monitoring.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213106.png"),
        "Figure 10. Feedback Review: confirmed error labels are retained for governed model evaluation.",
        alt_text="Feedback review page with confirmed feedback cases and error types.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213121.png"),
        "Figure 11. Model Monitoring: live operational metrics and benchmark-only evidence are visibly separated.",
        alt_text="Model monitoring page separating live operational metrics from research benchmark fidelity metrics.",
    )
    add_table(
        doc,
        ["CONTROL", "IMPLEMENTED OUTCOME"],
        [
            ["Feedback validation", "User opinion is not treated as validated until analyst confirmation"],
            ["Operational monitoring", "Scanned, quarantined, backlog, model error and feedback counts derive from stored workflow state"],
            ["Validation threshold", "The page says not enough confirmed feedback yet when live validation is insufficient"],
            ["Benchmark separation", "Fixed fidelity evidence is labelled as benchmark-only, not Gmail accuracy"],
        ],
        [2450, 6910],
    )

    doc.add_page_break()
    add_para(doc, "10. Dashboard Evidence: Model Readiness and Configuration", style="Heading 1")
    add_para(
        doc,
        "The system includes a readiness gate because classification is unsafe if live features do not match the "
        "training-time transformation. The Settings view exposes that gate and restricts model configuration to compatible teacher-surrogate pairs.",
    )
    add_figure(
        doc,
        screenshot("Screenshot 2026-05-25 213139.png"),
        "Figure 12. Settings: validated pipeline, controlled model pair and matched 292-feature model state.",
        alt_text="Settings page showing readiness status, Random Forest teacher, EBM surrogate and feature count 292.",
    )
    add_para(doc, "Recovered and Validated Model Pipeline", style="Heading 2")
    add_para(
        doc,
        "During implementation, the exact processed feature order was recovered from fitted EBM surrogate metadata. "
        "The fitted training preprocessor was deterministically reconstructed from the original training data and "
        "accepted only after reproducing the saved 292-feature processed split. The readiness panel exposes this state "
        "instead of silently trusting an incomplete feature vector.",
    )
    add_table(
        doc,
        ["READINESS ITEM", "EVIDENCE IN SYSTEM", "SAFETY PURPOSE"],
        [
            ["Teacher model", "Random Forest v1 Ready", "Confirms detector artifact is available"],
            ["Surrogate model", "EBM for Random Forest v1 Ready", "Confirms compatible explanation model"],
            ["Preprocessor", "Training preprocessor Ready", "Confirms transformation is available"],
            ["Feature order", "Recovered; feature count 292; Matched", "Prevents order/schema mismatch predictions"],
        ],
        [2300, 3100, 3960],
    )

    doc.add_page_break()
    add_para(doc, "11. Implementation and Safety Review", style="Heading 1")
    add_para(doc, "Technology and Data Boundaries", style="Heading 2")
    add_table(
        doc,
        ["AREA", "IMPLEMENTED APPROACH", "WHY IT MATTERS"],
        [
            ["Frontend", "React/Vite pages with reusable cards, badges, tables and status panels", "Presents investigation workflow cleanly"],
            ["Backend", "FastAPI endpoints and service orchestration", "Keeps mailbox/model operations behind validated APIs"],
            ["Database", "SQLite with email, prediction, explanation, feedback, sync-run and audit tables", "Supports traceability in the prototype"],
            ["Mailbox action", "Gmail label add/remove for quarantine and release", "No permanent deletion of messages"],
            ["Secrets", "Local environment values or ignored token files", "Avoids credential exposure in Git and logs"],
        ],
        [1800, 4160, 3400],
    )
    add_para(doc, "Safety Controls Already Demonstrated", style="Heading 2")
    controls = [
        ("Manual sync only", "The sidebar explicitly displays Manual Sync Only and exposes the Sync Gmail button."),
        ("Safe model readiness", "The UI identifies whether predictions can be trusted for quarantine actions."),
        ("Fail-safe design", "A model sync failure is designed to become model_error / needs_review rather than being silently allowed."),
        ("Reversible quarantine", "Containment is implemented with Gmail labels and an analyst Release path."),
        ("Feedback governance", "Confirmed feedback contributes to monitoring; automatic retraining is not enabled."),
        ("Auditability", "Quarantine, release, feedback review and model switching are recorded as audit events in the backend."),
    ]
    for label, text in controls:
        add_labeled_para(doc, label + ":", text)
    add_para(doc, "Verification Evidence", style="Heading 2")
    add_table(
        doc,
        ["CHECK", "MOST RECENT RESULT", "COVERAGE"],
        [
            ["Backend tests", "27 tests passed", "Readiness, audit, Gmail mapping, URL preservation, deduplication and workflow cases"],
            ["Frontend lint", "Passed", "JavaScript/React quality checks"],
            ["Frontend production build", "Passed", "Vite compilation of dashboard pages and components"],
        ],
        [2100, 2200, 5060],
    )

    doc.add_page_break()
    add_para(doc, "12. Honest Assessment and Current Limitations", style="Heading 1")
    add_para(
        doc,
        "PhishGuard is realistic as a university showcase and a technically meaningful prototype: it integrates a real "
        "Gmail workflow with an explainable ML decision-support interface, rather than showing only a notebook result or "
        "static dashboard. Its strongest contribution is the operational and governance design around the classifier.",
    )
    add_para(
        doc,
        "It should not yet be presented as a deployable enterprise phishing protection product. A real business rollout "
        "would demand stronger identity and access controls, deployment hardening, privacy retention controls, larger "
        "operational validation, incident handling, and richer email authentication and reputation signals.",
    )
    add_table(
        doc,
        ["CURRENT STRENGTH", "CURRENT LIMITATION / NEXT HARDENING STEP"],
        [
            ["Live Gmail manual workflow and reversible actions", "Add production authentication, authorisation and deployment controls"],
            ["Validated current 292-feature prediction pipeline", "Evaluate live performance with sufficient analyst-confirmed labels"],
            ["Local and global explanation interfaces", "Study explanation usefulness and analyst decision quality with real users"],
            ["Operational SQLite audit and status state", "Introduce stronger database/privacy retention and backup policy for deployment"],
            ["Controlled feedback review", "Keep retraining offline and validated; do not automate learning from user reports"],
            ["Existing engineered features", "Future research may assess authentication/reputation features after MVP stability"],
        ],
        [3500, 5860],
    )
    add_callout(
        doc,
        "My honest product position",
        "This is a credible applied software project because it converts ML research into a complete, safety-aware analyst workflow. "
        "Its value is not that it solves phishing by itself; its value is that it makes detection explainable, controllable and reviewable.",
        "blue",
    )

    doc.add_page_break()
    add_para(doc, "13. Showcase Demonstration Plan", style="Heading 1")
    add_para(
        doc,
        "A short demonstration should focus on the product workflow rather than presenting it as a model research experiment.",
    )
    add_table(
        doc,
        ["DEMO MOMENT", "WHAT TO SHOW", "MESSAGE TO THE AUDIENCE"],
        [
            ["1. Product opening", "Overview page and live/benchmark labels", "This is a Gmail operational dashboard with truthful reporting"],
            ["2. Readiness assurance", "Pipeline validated panel and Settings model pairing", "Trusted decisions require validated model artifacts"],
            ["3. Manual operation", "Press Sync Gmail and inspect sync status", "Analyst controls the mailbox scan and can see skipped/failed items"],
            ["4. Investigation", "Open a detection and show explanation factors", "The model gives evidence in readable and technical forms"],
            ["5. Reversible action", "Show Quarantine/Release workflow", "The prototype protects workflow continuity and never deletes mail"],
            ["6. Governance", "Feedback Review and Model Monitoring", "Errors become analyst-confirmed monitoring evidence, not automatic retraining"],
        ],
        [1850, 3370, 4140],
    )
    add_para(doc, "Feature Evidence Index", style="Heading 2")
    add_table(
        doc,
        ["FEATURE CLAIM", "SCREENSHOT EVIDENCE"],
        [
            ["Operational summary and manual sync status", "Figures 2-3"],
            ["Live Gmail detection queue and quarantine", "Figures 4-5"],
            ["Local explanation and analyst action workflow", "Figures 6-7"],
            ["Global model explanation and benchmark fidelity", "Figures 8-9"],
            ["Analyst feedback and truthful live monitoring", "Figures 10-11"],
            ["Validated model pairing and feature readiness", "Figure 12"],
        ],
        [4800, 4560],
    )
    add_para(doc, "Closing Statement", style="Heading 2")
    add_para(
        doc,
        "PhishGuard SOC demonstrates a realistic direction for applied cybersecurity engineering: a live Gmail prototype "
        "that joins phishing-risk classification to explanation, reversible response, feedback validation and monitoring. "
        "For a university showcase, it provides both technical depth and a clear business-security story, while maintaining "
        "honest boundaries about what remains necessary before real-world deployment.",
    )
    add_callout(
        doc,
        "Publication check before submission",
        "Create a redacted copy of this report if it will be distributed publicly, because the supporting screenshots contain live mailbox identifying information.",
        "red",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(path)
